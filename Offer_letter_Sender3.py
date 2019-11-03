import docx
import pandas as pd
from docx.shared import Pt
import os.path

from os import chdir, getcwd, listdir, path
from time import strftime
from win32com import client

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

from tkinter import *



def run_sender(file_nm):
    print(file_nm)
    # file_nm = input('Enter the filename(with extension): ')
    file = pd.read_csv(file_nm)
    names = file.iloc[:,0]
    emails = file.iloc[:,1]

    # print(emails)

    doc = docx.Document('Offer Letter - Campus Ambassador Program _ Aparoksha.docx')
    doc.paragraphs[6].text='Dear '
    doc.paragraphs[6].add_run()
    doc.paragraphs[6].runs[1].bold=True
    name = doc.paragraphs[6].runs[1]
    font = name.font
    font.size = Pt(14)

    for name in names:
        print(name)
        doc.paragraphs[6].runs[1].text=name +','
        doc.save(os.path.join("E:\\Projects\\OfferLetter_sender\\pdfs", (name+'.docx')))

    def count_files(filetype):
        count_files = 0
        for files in listdir(folder):
            if files.endswith(filetype):
                count_files += 1
        return count_files

    def check_path(prompt):
        abs_path = input(prompt)
        while path.exists(abs_path) != True:
            print ("\nThe specified path does not exist.\n")
            abs_path = input(prompt)
        return abs_path    
        
    print ("\n")

    folder = "E:\\Projects\\OfferLetter_sender\\pdfs"

    chdir(folder)


    num_docx = count_files(".docx")
    num_doc = count_files(".doc")


    if num_docx + num_doc == 0:
        print ("\nThe specified folder does not contain docx or docs files.\n")
        exit()
    else:
        print ("\nNumber of doc and docx files: ", num_docx + num_doc, "\n")
        print ("\n\nStarting to convert files ...\n")
        

    try:
        word = client.DispatchEx("Word.Application")
        for files in listdir(getcwd()):
            match = 0
            if files.endswith(".doc"): s, match = "doc", 1
            elif files.endswith(".docx"): s, match = "docx", 1
            if match:
                new_name = files.replace("."+s, r".pdf")
                in_file = path.abspath(folder + "\\" + files)
                new_file = path.abspath(folder + "\\" + new_name)
                doc = word.Documents.Open(in_file)
                print ('Conversion Completed (from .docx to .pdf) ', path.relpath(new_file))
                doc.SaveAs(new_file, FileFormat = 17)
                doc.Close()

    except (Exception, e):
        print (e)
    finally:
        word.Quit()

    print("\n", "Finished converting files to pdf format!!!")
    print("Starting to send email!!!")

    # Count the number of pdf files.

    # num_pdf = count_files(".pdf")   

    # print ("\nNumber of pdf files: ", num_pdf)

    # Check if the number of docx and doc file is equal to the number of files.

    # if num_docx + num_doc == num_pdf:
    #     print ("\nNumber of doc and docx files is equal to number of pdf files.")
    # else:
    #     print ("\nNumber of doc and docx files is not equal to number of pdf files.")






    os.chdir('E:\\Projects\\OfferLetter_sender\\pdfs')
    i=3
    j=0
    for name in names:
        filename = name + '.pdf'
        # print(os.getcwd)
        # print(filename)
        # print(type(filename))
        msg = MIMEMultipart()
        # msg['FROM']="biscuit01010000@gmail.com"
        msg['TO'] = "biscuit01010000@gmail.com"

        msg['Subject']='Hi This is an pdf for '+ name
        body = "Hello"
        msg.attach(MIMEText(body,'plain'))
        image = open(str(filename),'rb')

        part = MIMEBase('application','octet-stream')
        part.set_payload(image.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition','attachment;filename=' + filename)

        msg.attach(part)



        smtp=smtplib.SMTP('smtp.gmail.com',587)
        smtp.ehlo()
        smtp.starttls()
        smtp.ehlo()
        smtp.login('lokesh27dinu@gmail.com','Kars27032001')
        # subject='Hi I am Lokesh'
        # body='Hello! Welcome to Gmail2.'
        # message = f'Subject:{subject}\n\n{body}'
        smtp.sendmail('lokesh27dinu@gmail.com',emails[j],msg.as_string())
        smtp.close()
        print('DONE ' + 'for ' + name)
        j+=1
        # msg['Subject']='Hi This is an pdf'+ str(i)
        filename = None

run_sender('names.csv')
# window = Tk()


# file_info_shower = Label(window,text="Filename : ")
# file_nm_entered = Entry(window)
# file_nm = file_nm_entered.get()
# print(file_nm)
# btn = Button(window,text= "Submit",command =lambda: '')#run_sender(file_nm))

# file_info_shower.grid(row=0,column = 0)
# file_nm_entered.grid(row=0,column = 1)
# btn.grid(row = 1,column = 1)

# window.mainloop()



